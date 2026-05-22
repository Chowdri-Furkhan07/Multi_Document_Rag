from langchain_core.documents import Document
import pandas as pd
from bs4 import BeautifulSoup
import docx
import tempfile
import re
import requests
from pypdf import PdfReader

# ---------------- CLEAN TEXT ----------------
def clean_text(text):

    text = re.sub(r'\s+', ' ', str(text))
    text = re.sub(r'\n+', ' ', text)

    return text.strip()


# ---------------- SAFE PDF ----------------
def load_pdf(uploaded_file):

    documents = []

    try:

        with tempfile.NamedTemporaryFile(
            delete=False,
            suffix=".pdf"
        ) as tmp:

            tmp.write(uploaded_file.getbuffer())
            tmp_path = tmp.name

        reader = PdfReader(tmp_path)

        for page in reader.pages:

            text = page.extract_text()

            if text:

                documents.append(
                    Document(
                        page_content=clean_text(text)
                    )
                )

        return documents

    except Exception as e:

        print("PDF ERROR:", uploaded_file.name)
        print(e)

        return []


# ---------------- DOCX ----------------
def load_docx(file):

    try:

        doc = docx.Document(file)

        full_text = []

        for p in doc.paragraphs:

            if p.text.strip():

                full_text.append(p.text)

        combined_text = "\n".join(full_text)

        return [
            Document(
                page_content=clean_text(combined_text)
            )
        ]

    except Exception as e:

        print("DOCX ERROR:", file.name)
        print(e)

        return []


# ---------------- CSV ----------------
def load_csv(file):

    try:

        df = pd.read_csv(file)

        if len(df.columns) == 1:

            df = df[df.columns[0]].str.split(",", expand=True)

            df.columns = [
                "department",
                "policy",
                "value",
                "conditions",
                "remarks"
            ]

        df.columns = df.columns.str.strip().str.lower()

        documents = []

        for _, row in df.iterrows():

            text = f"""
Department: {row.get('department','NA')}
Policy: {row.get('policy','NA')}
Value: {row.get('value','NA')}
Conditions: {row.get('conditions','NA')}
Remarks: {row.get('remarks','NA')}
"""

            documents.append(
                Document(
                    page_content=clean_text(text)
                )
            )

        return documents

    except Exception as e:

        print("CSV ERROR:", file.name)
        print(e)

        return []


# ---------------- HTML / WEB ----------------
def load_web(file):

    try:

        if hasattr(file, "read"):

            html_content = file.read().decode("utf-8")

        else:

            response = requests.get(file)

            html_content = response.text

        soup = BeautifulSoup(html_content, "html.parser")

        for tag in soup(["script", "style"]):

            tag.decompose()

        text_parts = []

        for element in soup.find_all(
            ["h1", "h2", "h3", "p", "li"]
        ):

            content = element.get_text(
                separator=" ",
                strip=True
            )

            if content:

                text_parts.append(content)

        unique_lines = list(
            dict.fromkeys(text_parts)
        )

        final_text = "\n".join(unique_lines)

        return [
            Document(
                page_content=clean_text(final_text)
            )
        ]

    except Exception as e:

        print("HTML ERROR:", file.name)
        print(e)

        return []